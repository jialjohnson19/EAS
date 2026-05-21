from docx import Document
from src.models.project import Project

from docxtpl import DocxTemplate
from src.models.project import Project


def create_word_report(project: Project, eris_df, output_path):

    # Load your Word template
    doc = DocxTemplate("src/templates/phase1_template.docx")

    # -----------------------------
    # Build context from project
    # -----------------------------
    context = project.__dict__.copy()

    # -----------------------------
    # OPTIONAL: add ERIS-derived fields
    # -----------------------------
    context["eris_row_count"] = len(eris_df)

    # If you already built SQG/CESQG logic:
    # context["sqg_facility_text"] = ...
    # context["cesqg_facility_text"] = ...

    # -----------------------------
    # Render Word document
    # -----------------------------
    doc.render(context)

    # Save output
    doc.save(output_path)


# def create_word_report(project: Project, output_path):

#     doc = Document()

#     doc.add_heading("Phase I Environmental Site Assessment", level=0)

#     # -------------------------------------------------
#     # INTRODUCTION
#     # -------------------------------------------------

#     intro_paragraph = (
#         f"Environmental Assessment Services, Inc. (EAS) has completed a Phase I "
#         f"Environmental Site Assessment of the "
#         f"{project.property_type} property "
#         f"(subject property), located at "
#         f"{project.property_address} in "
#         f"{project.city_name}, "
#         f"{project.county_name} County, Florida. "
#         f"The {project.county_name} County Property Appraiser "
#         f"identifies the parcel as "
#         f"{project.parcel_number}. "
#         f"For the purposes of this report, the parcel will be referred to as "
#         f"the “subject property”. EAS personnel conducted the site reconnaissance "
#         f"of the subject property on "
#         f"{project.site_visit_date}. "
#         f"A brief legal description of the subject property from OR Book "
#         f"{project.book_number}, Page "
#         f"{project.page_number} is included in Appendix B "
#         f"of this report."
#     )

#     doc.add_paragraph(intro_paragraph)

#     # -------------------------------------------------
#     # SITE DESCRIPTION
#     # -------------------------------------------------

#     site_description = (
#         f"The subject property is located at {project.property_address}. "
#         f"The subject property is {project.property_size} acres. "
#         f"The subject property is currently occupied with/by "
#         f"{project.property_description}. "
#         f"There is {project.number_of_buildings} building(s) on the subject property. "
#         f"The building on the subject property is a "
#         f"{project.stories} story, "
#         f"{project.building_material} structure built in "
#         f"{project.year_built}. "
#         f"The building is approximately {project.gross_sqft} square feet, of which "
#         f"{project.climate_sqft} is climate-controlled area as calculated by the "
#         f"{project.county_name} County Property Appraiser. "
#         f"The remainder of the parcel is covered with "
#         f"{project.remainder_description}. "
#         f"Ingress and egress from the subject property is via "
#         f"{project.ingress_egress}."
#     )

#     doc.add_paragraph(site_description)

#     # -------------------------------------------------
#     # UTILITIES
#     # -------------------------------------------------

#     utilities_paragraph = (
#         f"Potable water and sewer services to the subject property are through "
#         f"the City of {project.city_name}. "
#         f"Trash is currently collected from the subject property in two dumpsters "
#         f"by the City of {project.city_name}. "
#         f"Electrical service is provided through {project.power_company}. "
#         f"There are {project.number_transformers} {project.transformer_type} electrical service "
#         f"transformers located {project.transformer_location} the subject property. "
#         f"The transformer(s) {project.pcb_label}. "
#         f"There were {project.natural_gas} natural gas connections to the building on the subject property. "
#         f"The storm water is directed to {project.stormwater}. "
#         f"There were {project.wells_present} monitor, irrigation or production wells "
#         f"located on the subject property that were noted by EAS personnel."
#     )

#     doc.add_paragraph(utilities_paragraph)

#     # -------------------------------------------------
#     # SURROUNDING PROPERTIES
#     # -------------------------------------------------

#     surrounding_properties = (
#         f"The subject property is located along the {project.general_location}. "
#         f"To the north of the subject property across {project.north_street} is {project.north_use}. "
#         f"To the south of the subject property is {project.south_use}. "
#         f"To the east of the subject property is {project.east_use}. "
#         f"To the west of the subject property is {project.west_use}. "
#         f"None of the adjacent properties appear to pose adverse environmental condition "
#         f"to the subject property."
#     )

#     doc.add_paragraph(surrounding_properties)

#     # -------------------------------------------------
#     # HISTORICAL REVIEW
#     # -------------------------------------------------

#     historical_review = (
#         f"An examination of available historic aerial photographs of the site vicinity "
#         f"was performed through the Florida Department of Transportation and Google Earth. "
#         f"The {project.aerial_dates} aerial photographs were reviewed in conjunction with this assessment. "
#         f"Although aerial photography does not exist for the subject property prior to "
#         f"{project.earliest_aerial}, it is not expected that this data gap interferes with the outcome "
#         f"of this assessment. "
#         f"An examination of available historic topographic maps was performed through USGS. "
#         f"The {project.topo_dates} maps were reviewed. "
#         f"The {project.sanborn_dates} were reviewed. "
#         f"City directories {project.city_directories} were reviewed. "
#         f"The subject property was historically undeveloped prior to construction in "
#         f"{project.year_built}. "
#         f"Tenants on the subject property have been {project.tenant_history}."
#     )

#     doc.add_paragraph(historical_review)

#     # -------------------------------------------------
#     # CONCLUSION
#     # -------------------------------------------------

#     conclusion_paragraph = (
#         "There were no current or historical recognized environmental conditions "
#         "associated with the subject property at the time of the site reconnaissance. "
#         "There are no known adverse historical environmental conditions on adjacent parcels. "
#         "There is no known migration of contaminated soil, groundwater, or vapors impacting the site."
#     )

#     doc.add_paragraph(conclusion_paragraph)

#     # -------------------------------------------------
#     # SAVE DOCUMENT
#     # -------------------------------------------------

#     doc.save(output_path)