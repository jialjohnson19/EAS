# report_writer.py
# This file writes the Phase I ESA report using ERIS data and manual inputs.

import os
import re

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt
from docx.oxml.ns import qn


# -------------------------------------------------
# HELPER FUNCTIONS
# -------------------------------------------------

def set_document_font(doc):
    """
    Sets the document font to Times New Roman, 12 pt.
    """

    styles = doc.styles

    for style in styles:
        if style.type == 1:  # paragraph styles
            font = style.font
            font.name = "Times New Roman"
            font.size = Pt(12)

            style.element.rPr.rFonts.set(
                qn("w:eastAsia"),
                "Times New Roman"
            )


def missing_input(key):
    """
    Returns placeholder text if user did not enter a value.
    Example: <ENTER CITY NAME>
    """

    return f"<ENTER {key.replace('_', ' ').upper()}>"


def add_highlighted_paragraph(doc, text):
    """
    Adds a paragraph and highlights missing input placeholders.
    Placeholders look like: <ENTER CITY NAME>
    """

    paragraph = doc.add_paragraph()

    parts = re.split(r"(<ENTER .*?>)", text)

    for part in parts:
        run = paragraph.add_run(part)

        if part.startswith("<ENTER") and part.endswith(">"):
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    return paragraph


def load_existing_paragraphs(existing_docx_path):
    """
    Reads an existing Word report and saves manually edited sections.

    This allows manual edits in the Word document to remain when the
    report is regenerated.
    """

    if not os.path.exists(existing_docx_path):
        return {}

    old_doc = Document(existing_docx_path)

    saved_sections = {}

    for paragraph in old_doc.paragraphs:
        text = paragraph.text.strip()

        if not text:
            continue

        if text.startswith("Environmental Assessment Services"):
            saved_sections["intro"] = text

        elif text.startswith("The subject property is located at"):
            saved_sections["site_description"] = text

        elif text.startswith("Potable water and sewer services"):
            saved_sections["utilities"] = text

        elif text.startswith("The subject property is located along"):
            saved_sections["surrounding_properties"] = text

        elif text.startswith("An examination of available historic aerial"):
            saved_sections["historical_review"] = text

        elif text.startswith("There were no current or historical recognized"):
            saved_sections["conclusion"] = text

    return saved_sections


# -------------------------------------------------
# MAIN REPORT FUNCTION
# -------------------------------------------------

def create_word_report(inputs, output_path):
    """
    Creates a Word report.

    First run:
        - Creates a baseline report.

    Later runs:
        - Reads the existing report.
        - Preserves manually edited sections.
        - Regenerates sections that have not been manually edited.
    """

    existing_sections = load_existing_paragraphs(output_path)

    doc = Document()
    set_document_font(doc)

    doc.add_heading("Phase I Environmental Site Assessment", level=0)

    # -------------------------------------------------
    # INTRODUCTION
    # -------------------------------------------------

    generated_intro = (
        f"Environmental Assessment Services, Inc. (EAS) has completed a Phase I "
        f"Environmental Site Assessment of the "
        f"{inputs.get('property_type') or missing_input('property_type')} property "
        f"(subject property), located at "
        f"{inputs.get('property_address') or missing_input('property_address')} in "
        f"{inputs.get('city_name') or missing_input('city_name')}, "
        f"{inputs.get('county_name') or missing_input('county_name')} County, Florida. "
        f"The {inputs.get('county_name') or missing_input('county_name')} County Property Appraiser "
        f"identifies the parcel as "
        f"{inputs.get('parcel_number') or missing_input('parcel_number')}. "
        f"For the purposes of this report, the parcel will be referred to as "
        f"the “subject property”. EAS personnel conducted the site reconnaissance "
        f"of the subject property on "
        f"{inputs.get('site_visit_date') or missing_input('site_visit_date')}. "
        f"A brief legal description of the subject property from OR Book "
        f"{inputs.get('book_number') or missing_input('book_number')}, Page "
        f"{inputs.get('page_number') or missing_input('page_number')} is included in Appendix B "
        f"of this report."
    )

    intro_paragraph = existing_sections.get("intro") or generated_intro
    add_highlighted_paragraph(doc, intro_paragraph)

    # -------------------------------------------------
    # SITE DESCRIPTION
    # -------------------------------------------------

    generated_site_description = (
        f"The subject property is located at "
        f"{inputs.get('property_address') or missing_input('property_address')}. "
        f"The subject property is approximately "
        f"{inputs.get('property_size') or missing_input('property_size')} acres. "
        f"The subject property is currently occupied with/by "
        f"{inputs.get('property_description') or missing_input('property_description')}. "
        f"There is "
        f"{inputs.get('number_of_buildings') or missing_input('number_of_buildings')} building(s) "
        f"on the subject property. "
        f"The building on the subject property is a "
        f"{inputs.get('stories') or missing_input('stories')} story, "
        f"{inputs.get('building_material') or missing_input('building_material')} structure built in "
        f"{inputs.get('year_built') or missing_input('year_built')}. "
        f"The building is approximately "
        f"{inputs.get('gross_sqft') or missing_input('gross_sqft')} square feet, of which "
        f"{inputs.get('climate_sqft') or missing_input('climate_sqft')} is climate-controlled area "
        f"as calculated by the "
        f"{inputs.get('county_name') or missing_input('county_name')} County Property Appraiser. "
        f"The remainder of the parcel is covered with "
        f"{inputs.get('remainder_description') or missing_input('remainder_description')}. "
        f"Ingress and egress from the subject property is via "
        f"{inputs.get('ingress_egress') or missing_input('ingress_egress')}."
    )

    site_description = existing_sections.get("site_description") or generated_site_description
    add_highlighted_paragraph(doc, site_description)

    # -------------------------------------------------
    # UTILITIES
    # -------------------------------------------------

    generated_utilities = (
        f"Potable water and sewer services to the subject property are through "
        f"the City of {inputs.get('city_name') or missing_input('city_name')}. "
        f"Trash is currently collected from the subject property in two dumpsters "
        f"by the City of {inputs.get('city_name') or missing_input('city_name')}. "
        f"Electrical service is provided through "
        f"{inputs.get('power_company') or missing_input('power_company')}. "
        f"There are "
        f"{inputs.get('number_transformers') or missing_input('number_transformers')} "
        f"{inputs.get('transformer_type') or missing_input('transformer_type')} electrical service "
        f"transformers located "
        f"{inputs.get('transformer_location') or missing_input('transformer_location')} the subject property. "
        f"The transformer(s) "
        f"{inputs.get('pcb_label') or missing_input('pcb_label')}. "
        f"There were "
        f"{inputs.get('natural_gas') or missing_input('natural_gas')} natural gas connections "
        f"to the building on the subject property. "
        f"The storm water is directed to "
        f"{inputs.get('stormwater') or missing_input('stormwater')}. "
        f"There were "
        f"{inputs.get('wells_present') or missing_input('wells_present')} monitor, irrigation or "
        f"production wells located on the subject property that were noted "
        f"by EAS personnel."
    )

    utilities_paragraph = existing_sections.get("utilities") or generated_utilities
    add_highlighted_paragraph(doc, utilities_paragraph)

    # -------------------------------------------------
    # SURROUNDING PROPERTIES
    # -------------------------------------------------

    generated_surrounding_properties = (
        f"The subject property is located along the "
        f"{inputs.get('general_location') or missing_input('general_location')}. "
        f"To the north of the subject property across "
        f"{inputs.get('north_street') or missing_input('north_street')} is "
        f"{inputs.get('north_use') or missing_input('north_use')}. "
        f"To the south of the subject property is "
        f"{inputs.get('south_use') or missing_input('south_use')}. "
        f"To the east of the subject property is "
        f"{inputs.get('east_use') or missing_input('east_use')}. "
        f"To the west of the subject property is "
        f"{inputs.get('west_use') or missing_input('west_use')}. "
        f"None of the adjacent properties appear to pose adverse environmental "
        f"condition to the subject property."
    )

    surrounding_properties = (
        existing_sections.get("surrounding_properties")
        or generated_surrounding_properties
    )
    add_highlighted_paragraph(doc, surrounding_properties)

    # -------------------------------------------------
    # HISTORICAL REVIEW
    # -------------------------------------------------

    generated_historical_review = (
        f"An examination of available historic aerial photographs of the site "
        f"vicinity was performed through the Florida Department of Transportation "
        f"and Google Earth. The "
        f"{inputs.get('aerial_dates') or missing_input('aerial_dates')} aerial photographs were reviewed "
        f"in conjunction with this environmental assessment. "
        f"Although aerial photography does not exist for the subject property "
        f"prior to {inputs.get('earliest_aerial') or missing_input('earliest_aerial')}, it is not expected "
        f"that this data gap interferes with the outcome of this assessment. "
        f"An examination of available historic topographic maps of the site vicinity "
        f"was performed through USGS. The "
        f"{inputs.get('topo_dates') or missing_input('topo_dates')} historic topographic maps were "
        f"reviewed in conjunction with this environmental assessment. "
        f"The {inputs.get('sanborn_dates') or missing_input('sanborn_dates')} were reviewed in conjunction "
        f"with this environmental assessment. "
        f"City Directories "
        f"{inputs.get('city_directories') or missing_input('city_directories')} reviewed in conjunction "
        f"with this environmental assessment. "
        f"According to resources reviewed, the subject property was historically "
        f"undeveloped prior to construction in "
        f"{inputs.get('year_built') or missing_input('year_built')}. "
        f"Tenants on the subject property have been "
        f"{inputs.get('tenant_history') or missing_input('tenant_history')}."
    )

    historical_review = (
        existing_sections.get("historical_review")
        or generated_historical_review
    )
    add_highlighted_paragraph(doc, historical_review)

    # -------------------------------------------------
    # CONCLUSION
    # -------------------------------------------------

    generated_conclusion = (
        "There were no current or historical recognized environmental conditions "
        "associated with the subject property at the time of the site reconnaissance. "
        "There are no known adverse historical environmental conditions on the parcels "
        "adjacent to the subject property. There is no known migration of contaminated "
        "soil, groundwater or vapors from contamination that have impacted the "
        "subject property."
        "\nEAS does not recommend additional environmental investigation of the subject property at this time."
    )

    conclusion_paragraph = existing_sections.get("conclusion") or generated_conclusion
    add_highlighted_paragraph(doc, conclusion_paragraph)

    # -------------------------------------------------
    # SAVE DOCUMENT
    # -------------------------------------------------

    doc.save(output_path)

    print(f"Report saved to: {output_path}")